from typing import Dict, Optional, BinaryIO, Tuple, List
from pathlib import Path
import docx
from PyPDF2 import PdfReader, PdfWriter
from striprtf.striprtf import rtf_to_text
from odf import text as odf_text
from odf.opendocument import load
from .pdf_handler import PDFHandler
from .language_detector import LanguageDetector

class DocumentHandler:
    def __init__(self):
        self.supported_formats = {
            '.txt': self._handle_txt,
            '.docx': self._handle_docx,
            '.pdf': self._handle_pdf,
            '.rtf': self._handle_rtf,
            '.odt': self._handle_odt
        }
        self.language_detector = LanguageDetector()
    
    def read_document(self, file_path: str, pages: Optional[List[int]] = None) -> Tuple[str, Optional[str]]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        text = self.supported_formats[extension](file_path, pages)
        clean_text = self._ensure_xml_compatible(text)
        detected_lang = self.language_detector.detect_language(clean_text)
        return clean_text, detected_lang

    def _ensure_xml_compatible(self, text: str) -> str:
        """Ensure text is XML compatible by removing invalid characters"""
        if not text:
            return ""
        # Remove NULL bytes and control characters while preserving newlines and tabs
        return ''.join(char for char in text if char in '\n\t' or (ord(char) >= 32 and ord(char) != 127))
    
    def _handle_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _handle_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def _handle_pdf(self, file_path: str, pages: Optional[List[int]] = None) -> str:
        pdf_handler = PDFHandler()
        text, layout_info = pdf_handler.extract_text_with_layout(file_path, pages)
        # Store layout info for later use when writing back to PDF
        self._current_pdf_layout = layout_info
        # Detect and store document structure
        self._current_pdf_structure = pdf_handler.detect_structure(file_path)
        return text
    
    def _handle_rtf(self, file_path: str) -> str:
        with open(file_path, 'r') as f:
            return rtf_to_text(f.read())
    
    def _handle_odt(self, file_path: str) -> str:
        doc = load(file_path)
        text_elements = doc.getElementsByType(odf_text.P)
        return '\n'.join([element.firstChild.data for element in text_elements if element.firstChild])

class DocumentWriter:
    def __init__(self):
        self.supported_formats = {
            '.txt': self._write_txt,
            '.docx': self._write_docx,
            '.pdf': self._write_pdf
        }
        self._pdf_handler = PDFHandler()
        self._current_pdf_layout = None
        self._current_pdf_structure = None
        self._current_target_lang = None

    def write_document(self, text: str, output_path: str, target_lang: str = None) -> None:
        path = Path(output_path)
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported output format: {extension}")
        
        self._current_target_lang = target_lang
        self.supported_formats[extension](text, output_path)
    
    def _write_txt(self, text: str, output_path: str) -> None:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    def _write_docx(self, text: str, output_path: str) -> None:
        doc = docx.Document()
        for paragraph in text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        doc.save(output_path)
    
    def _write_pdf(self, text: str, output_path: str) -> None:
        if not hasattr(self, '_current_pdf_layout'):
            # If no layout info available, create basic PDF
            writer = PdfWriter()
            page = writer.add_blank_page(width=612, height=792)
            # TODO: Add basic text rendering
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            return
        
        # Use layout-preserving PDF writing if layout info is available
        self._pdf_handler.write_pdf_with_layout(
            self._original_path,  # original path for image extraction
            text,
            output_path,
            self._current_pdf_layout,
            target_lang=self._current_target_lang
        )