import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from deep_translator import GoogleTranslator
import docx
import os
from PyPDF2 import PdfReader
from striprtf.striprtf import rtf_to_text
from odf import text as odf_text
from odf.opendocument import load
from ttkthemes import ThemedTk
import json
from pathlib import Path


class DocumentTranslator:

    # Add cancel flag as instance variable
    def __init__(self, root):
        self.root = root
        self.root.title("Smart Indian Language Document Translator")
        self.root.geometry("1000x800")
        self.root.set_theme("equilux")
        
        # Initialize settings
        self.settings_file = Path("translator_settings.json")
        self.load_settings()
        
        # Add translation cancel flag
        self.cancel_translation = False
        
        # Supported languages
        self.languages = {
            'Hindi': 'hi',
            'Bengali': 'bn',
            'Telugu': 'te',
            'Tamil': 'ta',
            'Marathi': 'mr',
            'Gujarati': 'gu',
            'Kannada': 'kn',
            'Malayalam': 'ml',
            'Punjabi': 'pa'
        }

        self.create_widgets()
        
    def load_settings(self):
        self.settings = {
            "theme": "dark",
            "recent_files": [],
            "custom_dictionary": {}
        }
        if self.settings_file.exists():
            try:
                with open(self.settings_file, 'r') as f:
                    self.settings.update(json.load(f))
            except:
                pass
    
    def save_settings(self):
        with open(self.settings_file, 'w') as f:
            json.dump(self.settings, f)
    
    def create_widgets(self):
        # Main container
        main_container = ttk.Frame(self.root, padding="10")
        main_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
        # Left panel for controls
        left_panel = ttk.Frame(main_container, padding="5")
        left_panel.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Right panel for preview
        right_panel = ttk.Frame(main_container, padding="5")
        right_panel.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_container.columnconfigure(1, weight=1)
        
        # File selection
        file_frame = ttk.LabelFrame(left_panel, text="Document Selection", padding="5")

        ttk.Label(file_frame, text="Select Document:").grid(row=0, column=0, sticky=tk.W)
        self.file_path = tk.StringVar()
        ttk.Entry(file_frame, textvariable=self.file_path, width=40).grid(row=0, column=1, padx=5)
        ttk.Button(file_frame, text="Browse", command=self.browse_file).grid(row=0, column=2)
        file_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # Recent files
        recent_frame = ttk.LabelFrame(left_panel, text="Recent Files", padding="5")
        self.recent_listbox = tk.Listbox(recent_frame, height=3)
        self.recent_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E))
        self.recent_listbox.bind('<Double-Button-1>', self.load_recent_file)
        recent_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)

        # Language and Format selection
        settings_frame = ttk.LabelFrame(left_panel, text="Translation Settings", padding="5")
        
        # Language selection
        ttk.Label(settings_frame, text="Target Language:").grid(row=0, column=0, sticky=tk.W)
        self.target_lang = tk.StringVar()
        lang_combo = ttk.Combobox(settings_frame, textvariable=self.target_lang,
                                 values=list(self.languages.keys()), width=20)
        lang_combo.grid(row=0, column=1, sticky=tk.W, padx=5)
        lang_combo.set("Hindi")
        
        # Output format selection
        ttk.Label(settings_frame, text="Output Format:").grid(row=1, column=0, sticky=tk.W, pady=(5,0))
        self.output_format = tk.StringVar(value="Same as Input")
        format_combo = ttk.Combobox(settings_frame, textvariable=self.output_format,
                                  values=["Same as Input", "DOCX", "PDF", "TXT", "RTF", "ODT"], width=20)
        format_combo.grid(row=1, column=1, sticky=tk.W, padx=5, pady=(5,0))
        
        settings_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # AI Instructions
        ai_frame = ttk.LabelFrame(left_panel, text="AI Translation Instructions", padding="5")
        self.ai_instructions = scrolledtext.ScrolledText(ai_frame, wrap=tk.WORD, width=40, height=4)
        self.ai_instructions.grid(row=0, column=0, sticky=(tk.W, tk.E))
        self.ai_instructions.insert(tk.END, "Enter any specific translation instructions or context here...")
        ai_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=5)

        # Translation controls
        control_frame = ttk.LabelFrame(left_panel, text="Translation Controls", padding="5")
        ttk.Button(control_frame, text="Translate Document",
                   command=self.translate_document).grid(row=0, column=0, pady=5)
        ttk.Button(control_frame, text="Preview Translation",
                   command=self.preview_translation).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(control_frame, text="Cancel",
                   command=self.cancel_translation_task).grid(row=0, column=2, padx=5, pady=5)
        ttk.Button(control_frame, text="Save As...",
                   command=self.save_translation_as).grid(row=0, column=3, padx=5, pady=5)
        control_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # Progress
        progress_frame = ttk.LabelFrame(left_panel, text="Progress", padding="5")
        self.progress_var = tk.StringVar()
        ttk.Label(progress_frame, textvariable=self.progress_var).grid(row=0, column=0)
        self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate')
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
        progress_frame.grid(row=4, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # Preview area
        preview_frame = ttk.LabelFrame(right_panel, text="Translation Preview", padding="5")
        self.preview_text = scrolledtext.ScrolledText(preview_frame, wrap=tk.WORD, width=50, height=30)
        self.preview_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        preview_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        right_panel.columnconfigure(0, weight=1)
        right_panel.rowconfigure(0, weight=1)

    def browse_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("PDF Files", "*.pdf"),
                ("RTF Files", "*.rtf"),
                ("ODT Files", "*.odt")
            ])
        if file_path:
            self.file_path.set(file_path)
            if file_path not in self.settings['recent_files']:
                self.settings['recent_files'].insert(0, file_path)
                if len(self.settings['recent_files']) > 5:
                    self.settings['recent_files'].pop()
                self.save_settings()
                self.update_recent_files()
            
    def load_recent_file(self, event):
        selection = self.recent_listbox.curselection()
        if selection:
            file_path = self.settings['recent_files'][selection[0]]
            self.file_path.set(file_path)
            
    def update_recent_files(self):
        self.recent_listbox.delete(0, tk.END)
        for file in self.settings['recent_files']:
            self.recent_listbox.insert(tk.END, Path(file).name)
            
    def preview_translation(self):
        input_file = self.file_path.get()
        if not input_file:
            messagebox.showerror("Error", "Please select a document first!")
            return
            
        try:
            # Get a sample of text to translate
            preview_text = ""
            if input_file.endswith('.docx'):
                doc = docx.Document(input_file)
                preview_text = "\n".join([p.text for p in doc.paragraphs[:3] if p.text.strip()])
            elif input_file.endswith('.txt'):
                with open(input_file, 'r', encoding='utf-8') as f:
                    preview_text = "\n".join(f.readlines()[:3])
            elif input_file.endswith('.pdf'):
                reader = PdfReader(input_file)
                if reader.pages:
                    preview_text = reader.pages[0].extract_text()
            elif input_file.endswith('.rtf'):
                with open(input_file, 'r') as f:
                    content = rtf_to_text(f.read())
                    preview_text = "\n".join(content.split('\n')[:3])
            elif input_file.endswith('.odt'):
                doc = load(input_file)
                elements = doc.getElementsByType(odf_text.P)
                preview_text = "\n".join([e.firstChild.data for e in elements[:3] if e.firstChild and e.firstChild.data.strip()])
                
            if preview_text:
                target = self.languages[self.target_lang.get()]
                translator = GoogleTranslator(source='auto', target=target)
                translated = translator.translate(text=preview_text)
                
                self.preview_text.delete(1.0, tk.END)
                self.preview_text.insert(tk.END, "Original Text:\n\n")
                self.preview_text.insert(tk.END, preview_text)
                self.preview_text.insert(tk.END, "\n\nTranslated Text:\n\n")
                self.preview_text.insert(tk.END, translated)
            else:
                messagebox.showwarning("Warning", "No text content found for preview")
                
        except Exception as e:
            messagebox.showerror("Error", f"Preview failed: {str(e)}")

    def cancel_translation_task(self):
        self.cancel_translation = True
        self.progress_var.set("Translation cancelled")
        self.progress_bar['value'] = 0
        self.root.update()

    def save_translation_as(self):
        input_file = self.file_path.get()
        if not input_file:
            messagebox.showerror("Error", "Please select a document first!")
            return

        output_format = self.output_format.get()
        if output_format == "Same as Input":
            output_format = os.path.splitext(input_file)[1][1:].upper()

        file_types = [
            ("Word Documents", "*.docx"),
            ("Text Files", "*.txt"),
            ("PDF Files", "*.pdf"),
            ("RTF Files", "*.rtf"),
            ("ODT Files", "*.odt")
        ]

        default_ext = f".{output_format.lower()}"
        suggested_name = f"{os.path.splitext(os.path.basename(input_file))[0]}_translated{default_ext}"

        save_path = filedialog.asksaveasfilename(
            defaultextension=default_ext,
            filetypes=file_types,
            initialfile=suggested_name
        )

        if save_path:
            self.translate_document(save_path)

    def translate_document(self, output_path=None):
        input_file = self.file_path.get()
        if not input_file:
            messagebox.showerror("Error", "Please select a document first!")
            return

        # Reset cancel flag
        self.cancel_translation = False
        
        target = self.languages[self.target_lang.get()]
        translator = GoogleTranslator(source='auto', target=target)
        max_retries = 3
        retry_delay = 2  # seconds
        
        # Get AI instructions
        ai_instructions = self.ai_instructions.get("1.0", tk.END).strip()
        if ai_instructions and ai_instructions != "Enter any specific translation instructions or context here...":
            self.preview_text.insert(tk.END, "\nAI Instructions:\n" + ai_instructions + "\n\n")
            
        # Determine output format and path
        output_format = self.output_format.get()
        if output_format == "Same as Input":
            output_format = os.path.splitext(input_file)[1][1:].upper()
        
        if not output_path:
            output_path = f"{os.path.splitext(input_file)[0]}_translated.{output_format.lower()}"
        
        # Reset progress bar and preview
        self.progress_bar['value'] = 0
        self.preview_text.delete(1.0, tk.END)
        translated = ""
        
        def translate_with_retry(text, retries=max_retries):
            if self.cancel_translation:
                return None
            for attempt in range(retries):
                try:
                    return translator.translate(text=text)
                except Exception as e:
                    if attempt < retries - 1 and not self.cancel_translation:
                        self.progress_var.set(f"Retrying translation... (Attempt {attempt + 2}/{retries})")
                        self.root.update()
                        import time
                        time.sleep(retry_delay)
                    else:
                        raise Exception(f"Translation failed after {retries} attempts: {str(e)}")
            return None

        try:
            if input_file.endswith('.docx'):
                doc = docx.Document(input_file)
                total_paragraphs = len(doc.paragraphs)

                for i, para in enumerate(doc.paragraphs, 1):
                    if self.cancel_translation:
                        break
                    if para.text.strip():
                        try:
                            translated = translate_with_retry(para.text)
                            if translated:
                                para.text = translated
                                self.preview_text.insert(tk.END, f"\n{translated}")
                            else:
                                if self.cancel_translation:
                                    break
                                raise Exception("Failed to translate paragraph")
                        except Exception as e:
                            if not self.cancel_translation:
                                self.preview_text.insert(tk.END, f"\nError translating paragraph: {str(e)}")
                            continue
                        progress = (i / total_paragraphs) * 100
                        self.progress_bar['value'] = progress
                        self.progress_var.set(f"Translating... {i}/{total_paragraphs} paragraphs")
                        self.root.update()

                if not self.cancel_translation:
                    if output_format == "DOCX":
                        doc.save(output_path)
                    else:
                        # Save to temporary DOCX first
                        temp_docx = f"{os.path.splitext(input_file)[0]}_temp.docx"
                        doc.save(temp_docx)
                        
                        try:
                            # Convert to desired format
                            if output_format == "PDF":
                                from docx2pdf import convert
                                convert(temp_docx, output_path)
                            elif output_format in ["TXT", "RTF", "ODT"]:
                                # Extract text and save in desired format
                                text_content = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                                
                                if output_format == "RTF":
                                    from striprtf.striprtf import rtf
                                    with open(output_path, 'w') as f:
                                        f.write(rtf(text_content))
                                elif output_format == "ODT":
                                    from odf.opendocument import OpenDocumentText
                                    from odf.text import P
                                    odt_doc = OpenDocumentText()
                                    for para in text_content.split('\n\n'):
                                        if para.strip():
                                            odt_doc.text.addElement(P(text=para))
                                    odt_doc.save(output_path)
                                else:  # TXT
                                    with open(output_path, 'w', encoding='utf-8') as f:
                                        f.write(text_content)
                        finally:
                            # Clean up temporary file
                            if os.path.exists(temp_docx):
                                os.remove(temp_docx)

            elif input_file.endswith('.txt'):
                with open(input_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                if not self.cancel_translation:
                    translated = translator.translate(text=content)
                    if translated:
                        if output_format == "TXT":
                            with open(output_path, 'w', encoding='utf-8') as f:
                                f.write(translated)
                        else:
                            # Create a temporary docx with the translated content
                            temp_docx = f"{os.path.splitext(input_file)[0]}_temp.docx"
                            doc = docx.Document()
                            for para in translated.split('\n'):
                                if para.strip():
                                    doc.add_paragraph(para)
                            doc.save(temp_docx)
                            
                            try:
                                # Convert to desired format
                                if output_format == "PDF":
                                    from docx2pdf import convert
                                    convert(temp_docx, output_path)
                                elif output_format == "RTF":
                                    from striprtf.striprtf import rtf
                                    text_content = '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                                    with open(output_path, 'w') as f:
                                        f.write(rtf(text_content))
                                elif output_format == "ODT":
                                    from odf.opendocument import OpenDocumentText
                                    from odf.text import P
                                    odt_doc = OpenDocumentText()
                                    for para in doc.paragraphs:
                                        if para.text.strip():
                                            odt_doc.text.addElement(P(text=para.text))
                                    odt_doc.save(output_path)
                            finally:
                                # Clean up temporary file
                                if os.path.exists(temp_docx):
                                    os.remove(temp_docx)
                        
                        self.preview_text.insert(tk.END, f"\n{translated}")

            elif input_file.endswith('.pdf'):
                reader = PdfReader(input_file)
                content = []
                total_pages = len(reader.pages)

                for i, page in enumerate(reader.pages, 1):
                    if self.cancel_translation:
                        break
                    text = page.extract_text()
                    if text.strip():
                        try:
                            translated = translate_with_retry(text)
                            if translated:
                                content.append(translated)
                                self.preview_text.insert(tk.END, f"\n{translated}")
                            else:
                                if self.cancel_translation:
                                    break
                                raise Exception("Failed to translate page")
                        except Exception as e:
                            if not self.cancel_translation:
                                self.preview_text.insert(tk.END, f"\nError translating page {i}: {str(e)}")
                            continue
                    progress = (i / total_pages) * 100
                    self.progress_bar['value'] = progress
                    self.progress_var.set(f"Translating... {i}/{total_pages} pages")
                    self.root.update()

                if not self.cancel_translation:
                    translated_content = '\n\n'.join(content)
                    
                    if output_format == "TXT":
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(translated_content)
                    else:
                        # Create a temporary docx with the translated content
                        temp_docx = f"{os.path.splitext(input_file)[0]}_temp.docx"
                        doc = docx.Document()
                        for para in translated_content.split('\n\n'):
                            if para.strip():
                                doc.add_paragraph(para)
                        doc.save(temp_docx)
                        
                        try:
                            # Convert to desired format
                            if output_format == "PDF":
                                from docx2pdf import convert
                                convert(temp_docx, output_path)
                            elif output_format == "RTF":
                                from striprtf.striprtf import rtf
                                text_content = '\n\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                                with open(output_path, 'w') as f:
                                    f.write(rtf(text_content))
                            elif output_format == "ODT":
                                from odf.opendocument import OpenDocumentText
                                from odf.text import P
                                odt_doc = OpenDocumentText()
                                for para in doc.paragraphs:
                                    if para.text.strip():
                                        odt_doc.text.addElement(P(text=para.text))
                                odt_doc.save(output_path)
                        finally:
                            # Clean up temporary file
                            if os.path.exists(temp_docx):
                                os.remove(temp_docx)

            elif input_file.endswith('.rtf'):
                with open(input_file, 'r') as f:
                    content = rtf_to_text(f.read())
                if not self.cancel_translation:
                    translated = translator.translate(text=content)
                    if translated:
                        if output_format == "RTF":
                            with open(output_path, 'w') as f:
                                f.write(rtf(translated))
                        else:
                            # Create a temporary docx with the translated content
                            temp_docx = f"{os.path.splitext(input_file)[0]}_temp.docx"
                            doc = docx.Document()
                            for para in translated.split('\n'):
                                if para.strip():
                                    doc.add_paragraph(para)
                            doc.save(temp_docx)
                            
                            try:
                                # Convert to desired format
                                if output_format == "PDF":
                                    from docx2pdf import convert
                                    convert(temp_docx, output_path)
                                elif output_format == "TXT":
                                    with open(output_path, 'w', encoding='utf-8') as f:
                                        f.write(translated)
                                elif output_format == "ODT":
                                    from odf.opendocument import OpenDocumentText
                                    from odf.text import P
                                    odt_doc = OpenDocumentText()
                                    for para in doc.paragraphs:
                                        if para.text.strip():
                                            odt_doc.text.addElement(P(text=para.text))
                                    odt_doc.save(output_path)
                            finally:
                                # Clean up temporary file
                                if os.path.exists(temp_docx):
                                    os.remove(temp_docx)
                        
                        self.preview_text.insert(tk.END, f"\n{translated}")

            elif input_file.endswith('.odt'):
                doc = load(input_file)
                elements = doc.getElementsByType(odf_text.P)
                content = "\n\n".join([e.firstChild.data for e in elements if e.firstChild and e.firstChild.data.strip()])
                
                if not self.cancel_translation:
                    translated = translator.translate(text=content)
                    if translated:
                        if output_format == "ODT":
                            odt_doc = OpenDocumentText()
                            for para in translated.split('\n\n'):
                                if para.strip():
                                    odt_doc.text.addElement(P(text=para))
                            odt_doc.save(output_path)
                        else:
                            # Create a temporary docx with the translated content
                            temp_docx = f"{os.path.splitext(input_file)[0]}_temp.docx"
                            doc = docx.Document()
                            for para in translated.split('\n\n'):
                                if para.strip():
                                    doc.add_paragraph(para)
                            doc.save(temp_docx)
                            
                            try:
                                # Convert to desired format
                                if output_format == "PDF":
                                    from docx2pdf import convert
                                    convert(temp_docx, output_path)
                                elif output_format == "TXT":
                                    with open(output_path, 'w', encoding='utf-8') as f:
                                        f.write(translated)
                                elif output_format == "RTF":
                                    from striprtf.striprtf import rtf
                                    with open(output_path, 'w') as f:
                                        f.write(rtf(translated))
                            finally:
                                # Clean up temporary file
                                if os.path.exists(temp_docx):
                                    os.remove(temp_docx)
                        
                        self.preview_text.insert(tk.END, f"\n{translated}")

            if not self.cancel_translation:
                self.progress_var.set("Translation completed")
                self.progress_bar['value'] = 100
                messagebox.showinfo("Success", "Translation completed successfully!")
            else:
                self.progress_var.set("Translation cancelled")
                self.progress_bar['value'] = 0
                messagebox.showinfo("Cancelled", "Translation was cancelled by user")

        except Exception as e:
            self.progress_var.set("Translation failed")
            self.progress_bar['value'] = 0
            messagebox.showerror("Error", f"Translation failed: {str(e)}")


if __name__ == '__main__':
    root = ThemedTk()
    app = DocumentTranslator(root)
    root.mainloop()
